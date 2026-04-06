"""Export transcript and summary artifacts."""

from __future__ import annotations

import json
import sys
from datetime import datetime, timezone
from html import escape
from pathlib import Path

from diaricat.errors import DiaricatError, ErrorCode
from diaricat.models.domain import ExportFormat, Project, SummaryDocument, TranscriptDocument
from diaricat.settings import Settings
from diaricat.utils.paths import project_exports_dir
from diaricat.utils.validation import sec_to_timestamp


class ExportService:
    BRAND_NAME = "Diarcat"
    BRAND_COLOR_HEX = "7F56D9"

    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    @staticmethod
    def _runtime_candidates() -> list[Path]:
        project_root = Path(__file__).resolve().parents[3]
        bundled_root = Path(getattr(sys, "_MEIPASS", project_root))
        return [
            bundled_root / "frontend_dist" / "assets",
            project_root / "frontend" / "src" / "assets",
            project_root / "assets",
        ]

    def _resolve_logo_path(self) -> Path | None:
        for base in self._runtime_candidates():
            if not base.exists():
                continue
            for pattern in ("diarcat-logo*.png", "diaricat-logo*.png"):
                matches = sorted(base.glob(pattern))
                if matches:
                    return matches[0]
        return None

    @staticmethod
    def _segment_line(segment, include_timestamps: bool) -> str:
        label = segment.speaker_name or segment.speaker_id
        text = segment.text_corrected or segment.text_raw
        if include_timestamps:
            return f"[{sec_to_timestamp(segment.start)} - {sec_to_timestamp(segment.end)}] {label}: {text}"
        return f"{label}: {text}"

    def _write_json(
        self,
        path: Path,
        project: Project,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
    ) -> None:
        payload = {
            "project": project.model_dump(mode="json"),
            "transcript": transcript.model_dump(mode="json"),
            "summary": summary.model_dump(mode="json") if summary else None,
        }
        path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")

    def _write_md(
        self,
        path: Path,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
        include_timestamps: bool,
    ) -> None:
        lines: list[str] = ["# Transcript", ""]
        for segment in transcript.segments:
            label = segment.speaker_name or segment.speaker_id
            text = segment.text_corrected or segment.text_raw
            if include_timestamps:
                lines.append(
                    f"- [{sec_to_timestamp(segment.start)} - {sec_to_timestamp(segment.end)}] **{label}**: {text}"
                )
            else:
                lines.append(f"- **{label}**: {text}")

        if summary:
            lines.extend(["", "# Summary", "", f"## Overview\n{summary.overview}"])
            lines.append("\n## Key Points")
            for point in summary.key_points:
                lines.append(f"- {point}")
            lines.append("\n## Decisions")
            for decision in summary.decisions:
                lines.append(f"- {decision}")
            lines.append("\n## Topics")
            for topic in summary.topics:
                lines.append(f"- {topic}")

        path.write_text("\n".join(lines), encoding="utf-8")

    def _write_txt(
        self,
        path: Path,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
        include_timestamps: bool,
    ) -> None:
        lines = [self._segment_line(segment, include_timestamps) for segment in transcript.segments]

        if summary:
            lines.extend(
                [
                    "",
                    "SUMMARY",
                    f"Overview: {summary.overview}",
                    "Key points:",
                    *[f"- {p}" for p in summary.key_points],
                    "Decisions:",
                    *[f"- {d}" for d in summary.decisions],
                    "Topics:",
                    *[f"- {t}" for t in summary.topics],
                ]
            )

        path.write_text("\n".join(lines), encoding="utf-8")

    def _write_docx(
        self,
        path: Path,
        project: Project,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
        include_timestamps: bool,
    ) -> None:
        try:
            from docx import Document  # type: ignore
            from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
            from docx.shared import Inches, Pt, RGBColor  # type: ignore
        except Exception as exc:
            raise DiaricatError(
                ErrorCode.EXPORT_ERROR,
                "Missing dependency for DOCX export (python-docx).",
                details=str(exc),
            ) from exc

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        logo_path = self._resolve_logo_path()
        if logo_path and logo_path.exists():
            header_logo = doc.add_paragraph()
            header_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_logo.add_run()
            run.add_picture(str(logo_path), width=Inches(1.0))

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{self.BRAND_NAME} - Transcripcion Inteligente")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor.from_string(self.BRAND_COLOR_HEX)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Documento de transcripcion y resumen")
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(90, 90, 90)

        metadata = doc.add_paragraph()
        metadata.paragraph_format.space_before = Pt(8)
        metadata.paragraph_format.space_after = Pt(12)
        metadata.style = doc.styles["Normal"]
        metadata_run = metadata.add_run(
            f"Proyecto: {project.id}\n"
            f"Origen: {project.source_path}\n"
            f"Generado: {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"
        )
        metadata_run.font.size = Pt(9)

        heading = doc.add_paragraph()
        heading_run = heading.add_run("Transcripcion")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        heading_run.font.color.rgb = RGBColor.from_string(self.BRAND_COLOR_HEX)

        for segment in transcript.segments:
            line = self._segment_line(segment, include_timestamps)
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.name = "Calibri"

        if summary:
            summary_title = doc.add_paragraph()
            summary_run = summary_title.add_run("Resumen ejecutivo")
            summary_run.bold = True
            summary_run.font.size = Pt(14)
            summary_run.font.color.rgb = RGBColor.from_string(self.BRAND_COLOR_HEX)
            summary_title.paragraph_format.space_before = Pt(14)

            doc.add_paragraph(f"Vision general: {summary.overview}")

            kp_title = doc.add_paragraph("Puntos clave")
            kp_title.runs[0].bold = True
            for point in summary.key_points:
                doc.add_paragraph(point, style="List Bullet")

            dec_title = doc.add_paragraph("Decisiones")
            dec_title.runs[0].bold = True
            for decision in summary.decisions:
                doc.add_paragraph(decision, style="List Bullet")

            topics_title = doc.add_paragraph("Temas")
            topics_title.runs[0].bold = True
            for topic in summary.topics:
                doc.add_paragraph(topic, style="List Bullet")

        doc.save(path)

    def _write_pdf(
        self,
        path: Path,
        project: Project,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
        include_timestamps: bool,
    ) -> None:
        try:
            from reportlab.lib import colors  # type: ignore
            from reportlab.lib.pagesizes import A4  # type: ignore
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore
            from reportlab.lib.units import cm  # type: ignore
            from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer  # type: ignore
        except Exception as exc:
            raise DiaricatError(
                ErrorCode.EXPORT_ERROR,
                "Missing dependency for PDF export (reportlab).",
                details=str(exc),
            ) from exc

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.8 * cm,
            title=f"{self.BRAND_NAME} Transcript",
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DiarcatTitle",
            parent=styles["Heading1"],
            textColor=colors.HexColor("#7F56D9"),
            fontSize=18,
            leading=22,
            spaceAfter=4,
        )
        h2_style = ParagraphStyle(
            "DiarcatH2",
            parent=styles["Heading2"],
            textColor=colors.HexColor("#7F56D9"),
            fontSize=13,
            leading=17,
            spaceBefore=8,
            spaceAfter=5,
        )
        normal_style = ParagraphStyle(
            "DiarcatNormal",
            parent=styles["BodyText"],
            fontSize=10,
            leading=14,
        )
        muted_style = ParagraphStyle(
            "DiarcatMuted",
            parent=styles["BodyText"],
            textColor=colors.HexColor("#555555"),
            fontSize=8.8,
            leading=12,
            spaceAfter=2,
        )

        elements = []
        logo_path = self._resolve_logo_path()
        if logo_path and logo_path.exists():
            elements.append(Image(str(logo_path), width=2.2 * cm, height=2.2 * cm))
            elements.append(Spacer(1, 0.2 * cm))

        elements.append(Paragraph(f"{self.BRAND_NAME} - Transcripcion Inteligente", title_style))
        elements.append(Paragraph("Documento de transcripcion y resumen", muted_style))
        elements.append(
            Paragraph(
                escape(
                    f"Proyecto: {project.id} | "
                    f"Origen: {project.source_path} | "
                    f"Generado: {datetime.now(tz=timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}"
                ),
                muted_style,
            )
        )
        elements.append(Spacer(1, 0.25 * cm))

        elements.append(Paragraph("Transcripcion", h2_style))
        for segment in transcript.segments:
            elements.append(Paragraph(escape(self._segment_line(segment, include_timestamps)), normal_style))
            elements.append(Spacer(1, 0.12 * cm))

        if summary:
            elements.append(Spacer(1, 0.3 * cm))
            elements.append(Paragraph("Resumen ejecutivo", h2_style))
            elements.append(Paragraph(escape(f"Vision general: {summary.overview}"), normal_style))
            elements.append(Spacer(1, 0.15 * cm))

            elements.append(Paragraph("Puntos clave", h2_style))
            for point in summary.key_points:
                elements.append(Paragraph(escape(f"- {point}"), normal_style))

            elements.append(Paragraph("Decisiones", h2_style))
            for decision in summary.decisions:
                elements.append(Paragraph(escape(f"- {decision}"), normal_style))

            elements.append(Paragraph("Temas", h2_style))
            for topic in summary.topics:
                elements.append(Paragraph(escape(f"- {topic}"), normal_style))

        doc.build(elements)

    def export(
        self,
        project: Project,
        transcript: TranscriptDocument,
        summary: SummaryDocument | None,
        formats: list[ExportFormat],
        include_timestamps: bool,
    ) -> dict[str, str]:
        exports_dir = project_exports_dir(self.settings, project.id)
        artifacts: dict[str, str] = {}

        for fmt in formats:
            if fmt == ExportFormat.JSON:
                path = exports_dir / "result.json"
                self._write_json(path, project, transcript, summary)
                artifacts[fmt.value] = str(path)

            elif fmt == ExportFormat.MD:
                path = exports_dir / "result.md"
                self._write_md(path, transcript, summary, include_timestamps)
                artifacts[fmt.value] = str(path)

            elif fmt == ExportFormat.TXT:
                path = exports_dir / "result.txt"
                self._write_txt(path, transcript, summary, include_timestamps)
                artifacts[fmt.value] = str(path)

            elif fmt == ExportFormat.PDF:
                path = exports_dir / "result.pdf"
                self._write_pdf(path, project, transcript, summary, include_timestamps)
                artifacts[fmt.value] = str(path)

            elif fmt == ExportFormat.DOCX:
                path = exports_dir / "result.docx"
                self._write_docx(path, project, transcript, summary, include_timestamps)
                artifacts[fmt.value] = str(path)

        return artifacts
